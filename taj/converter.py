import os

import docx

from builder import load_file


def get_words(text):
    return " ".join(text.split(' ')[1:]).rstrip()


class Converter(object):
    def add_hyperlink(self, paragraph, text, url):
        """
        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)
        return hyperlink

    def convert(self, type, online_folder, chunks_text_path, output_folder):
        document = docx.Document()
        with open(f'{chunks_text_path}/text', 'r') as texts, open(f'{chunks_text_path}/segments') as segments, open(f'{chunks_text_path}/wav.scp') as wavs:
            for index, (text, segment, wav) in enumerate(zip(texts, segments, wavs)):
                text_words = get_words(text)
                text_utt = text.split(' ')[0].rstrip()
                seg_utt = segment.split(' ')[0].rstrip()
                seg_name = segment.split(' ')[1].rstrip()
                wav_name = wav.split(' ')[0].rstrip()
                wav_file = wav.split(' ')[1].rstrip()
                file_name = wav_file.split("/")[-1].rstrip()
                if text_utt != seg_utt:
                    raise NameError(f'utterance: {text_utt}/{seg_utt} - not matched')
                if seg_name != wav_name:
                    raise NameError(f'name: {seg_name}/{wav_name} - not matched')
                p = document.add_paragraph("")
                # print(f'({index}) ')
                print(f'{text_words}')
                print(f'{online_folder}/{file_name}')
                self.add_hyperlink(p, text_words, f'{online_folder}/{file_name}')
        output_file_path = f'{output_folder}/{chunks_text_path.split("/")[-1]}.doc'
        if not os.path.exists(output_folder):
            os.mkdir(output_folder)
        document.save(output_file_path)
        return output_file_path

