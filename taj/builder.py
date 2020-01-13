import shlex
import subprocess
import re
import json
import os
import docx
from exceptions import InputError


# def add_hyperlink(paragraph, text, url):
#     """
#     :param paragraph: The paragraph we are adding the hyperlink to.
#     :param url: A string containing the required url
#     :param text: The text displayed for the url
#     :return: The hyperlink object
#     """
#     # This gets access to the document.xml.rels file and gets a new relation id value
#     part = paragraph.part
#     r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
#     # Create the w:hyperlink tag and add needed values
#     hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
#     hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
#     # Create a w:r element and a new w:rPr element
#     new_run = docx.oxml.shared.OxmlElement('w:r')
#     rPr = docx.oxml.shared.OxmlElement('w:rPr')
#     # Join all the xml elements together add add the required text to the w:r element
#     new_run.append(rPr)
#     new_run.text = text
#     hyperlink.append(new_run)
#     # Create a new Run object and add the hyperlink into it
#     r = paragraph.add_run()
#     r._r.append(hyperlink)
#     return hyperlink


class ChunkBuilder(object):
    # def compose(self, mode, transcript, audio_source,
    #             markup_file, split_sentences, audio_output,
    #             transcription_output, validate, audio_output_chunks,
    #             word_output_file, online_folder, doc_output):
    #     if mode == 'make_markup':
    #         self.make_markup(transcript, markup_file, split_sentences)
    #     if mode == 'chunk':
    #         self.chunk(audio_source, transcript, markup_file, audio_output)
    #     if mode == 'transcribe':
    #         self.transcribe_audio(audio_source, doc_output)
    #     if mode == 'word':
    #         self.word(audio_output_chunks, word_output_file, online_folder, markup_file)
    #     if mode == 'batch':
    #         self.batch(audio_source, word_output_file, transcription_output, online_folder,
    #                    markup_file, split_sentences, audio_output, doc_output)

    def build(self, ffmpeg_cli):
        for ffmpeg_cli_item in ffmpeg_cli:
            options = shlex.split(ffmpeg_cli_item)
            subprocess.call(options)

    def batch(self, audio_source, word_output_file, transcription_output, online_folder, markup_file, split_sentences,
              audio_output, doc_output):
        print('Start transcribe...')
        transcription_path = self.transcribe_audio(audio_source, transcription_output)
        print(f'Start Markup...{transcription_path}')
        markup_path = self.make_markup(transcription_path, markup_file, split_sentences)
        print(f'Start Making chunk phrases...{markup_path}')
        markup_file = load_file(markup_path)
        transcription_file = load_json(transcription_path)
        chunk_phrase = self.compile_chunk_phrases(transcription_file, markup_file, audio_output)
        print(f'Start ffmpeg making cli lines...{chunk_phrase}')
        ffmpeg_cli = self.compile_ffmpeg_cli(chunk_phrase, audio_source)
        print(f'Start ffmpeg making chunks...{ffmpeg_cli}')
        self.build(ffmpeg_cli)
        print('Start Making word file...')
        self.word(audio_output, word_output_file, online_folder, markup_path)
        print(f'Upload chunks to the online folder: {online_folder}')
        print(f'Word file can be found here: {word_output_file}')
        return word_output_file

    # def word(self, audio_output_chunks, word_output_file, online_folder, markup_file):
    #     document = docx.Document()
    #     name_of_chunk = audio_output_chunks.split("/")[-1]
    #     markup_file_text = load_file(markup_file)
    #     phrase_list = []
    #     for phrase in markup_file_text.split("|"):
    #         if len(phrase) > 0:
    #             phrase_list.append(phrase)
    #     for index, phrase in enumerate(phrase_list):
    #         p = document.add_paragraph(phrase)
    #         add_hyperlink(p, f'[{index + 1}]  ', f'{online_folder}{name_of_chunk}{index + 1}.mp3')
    #     document.save(word_output_file)
    #     return word_output_file

    def make_markup(self, transcription_path, path_man, path_auto):
        transcript = load_json(transcription_path)
        if path_auto == '':
            path = path_man
        else:
            path = self.make_markup_file(transcript, path_auto)
        return path

    def chunk(self, audio_path, transcript_path, markup_path, audio_output_path):
        transcript = load_json(transcript_path)
        markup_file = load_file(markup_path)
        chunk_phrases = self.compile_chunk_phrases(transcript, markup_file, audio_output_path)
        ffmpeg_cli_list = self.compile_ffmpeg_cli(chunk_phrases, audio_path)
        list_of_output_files = []
        for chunk_phrase in chunk_phrases:
            list_of_output_files.append(f'{chunk_phrase["name"]}.mp3')
        self.build(ffmpeg_cli_list)
        return list_of_output_files

    def compile_ffmpeg_cli(self, chunk_phrases, audiofile):
        ffmpeg_cli = []
        if len(chunk_phrases) == 0:
            raise InputError('Something is wrong - chunk phrase list empty!')
        if audiofile.split(".")[-1] == 'wav':
            audiofile = self.convert_wav_to_mp3(audiofile)
        if audiofile.split(".")[-1] != 'mp3':
            raise Exception("OMG! The input audio file needs to be in wav or mp3 format.")

        for chunk_phrase in chunk_phrases:
            ffmpeg_cli.append(
                f'ffmpeg -i {audiofile} -ss {chunk_phrase["start"]} -to {chunk_phrase["end"]} -c copy -y {chunk_phrase["name"]}.mp3')

        ffmpeg_cli[
            -1] = f'ffmpeg -i {audiofile} -ss {chunk_phrases[-1]["start"]} -c copy -y {chunk_phrases[-1]["name"]}.mp3'
        return ffmpeg_cli

    def compile_chunk_phrases(self, transcription, markedup_taj_file, audio_output_path):
        chunk_list = markedup_taj_file.split('|')
        chunk_word_list = []
        words = transcription['words']
        word_count_end = 0
        chunk_phrases = []
        for chunk in chunk_list:
            chunk_words = chunk.split()
            if len(chunk_words) > 0:
                chunk_word_list.append(chunk_words)
        for index, chunk_word in enumerate(chunk_word_list):
            word_count_start = word_count_end
            word_count_end = word_count_start + len(chunk_word)
            start = words[word_count_start]['start']
            end = words[word_count_end - 1]['end']
            chunk_phrases.append({'name': f'{audio_output_path}{index + 1}', 'start': start, 'end': end})
        return chunk_phrases

    def transcribe_audio(self, audio_source, doc_output):
        current_working_dir = os.getcwd()
        audio_file_name = self.make_file_name(audio_source)
        instruction = f'docker run --rm  -v "{current_working_dir}:/tmp/media" --name bbc-kaldi-container  artifactory-noforge.virt.ch.bbc.co.uk:8443/bbc-kaldi:0.0.11 bbc-kaldi /tmp/media/{audio_source} /tmp/media/{doc_output}'
        transcription_output_path = f'{doc_output}/results/{audio_file_name}/transcription.json'
        options = shlex.split(instruction)
        print(f'{options}')
        if not os.path.exists(transcription_output_path):
            print(f'Making transcript....................................................................')
            subprocess.call(options)
        return transcription_output_path

    def make_file_name(self, audio_source):
        audio_file_name_full = audio_source.split("/")[-1]
        audio_file_name = ".".join(audio_file_name_full.split(".")[:-1])
        return audio_file_name

    def make_markup_file(self, transcription, path):
        punct = transcription["punct"]
        markup = re.sub(r"\.", ".|", punct)
        with open(path, "w", encoding='utf-8') as fp:
            fp.write(markup)
        return path

    def extract_markup_file(self, transcription, path):
        with open(path, "w", encoding='utf-8') as fp:
            fp.write(transcription['punct'])
        return path

    def convert_wav_to_mp3(self, wav_file):
        mp3_file = f'{wav_file.split(".")[0]}.mp3'
        options = [
            'ffmpeg',
            '-y',
            '-i', wav_file,
            '-vn',
            '-ar', '44100',
            '-ac', '2',
            '-b:a',
            '192k',
            mp3_file
        ]
        subprocess.call(options)
        return f'{wav_file.split(".")[0]}.mp3'


def load_json(path):
    with open(path, encoding='utf-8') as fp:
        return json.load(fp)


def load_file(path):
    with open(path, encoding='utf-8') as fp:
        return fp.read()

    # ffmpeg -i 20191130-2034_Test1.wav -vn -ar 44100 -ac 2 -b:a 192k test1.mp3
    # ffmpeg -i fixtures/20191130-2034_Test1.wav -ss 0.1 -to 6.26 -c copy chunk1.mp3
